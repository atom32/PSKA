from __future__ import annotations

from dataclasses import dataclass, field
from fnmatch import fnmatch
import hashlib
import io
import json
import mimetypes
from pathlib import Path
import tempfile
from typing import Any
from uuid import NAMESPACE_URL, uuid5
import zipfile
from xml.etree import ElementTree

from pska_core.config import DEFAULT_FILES_MAX_BYTES, DEFAULT_SPREADSHEET_MAX_COLUMNS, DEFAULT_SPREADSHEET_MAX_ROWS_PER_SHEET
from pska_core.connectors import connector_record_to_payload, connector_state_from_mapping, connector_state_id
from pska_core.enums import Visibility
from pska_core.ingest import IngestService
from pska_core.models import DEFAULT_TENANT_ID, Chunk, ConnectorState, Document, SourceItem
from pska_core.offline_index import OfflineIndexService
from pska_core.processing import resolve_processing_config
from pska_core.store import KnowledgeStore


TEXT_SUFFIXES = {
    ".cfg",
    ".conf",
    ".csv",
    ".json",
    ".jsonl",
    ".log",
    ".md",
    ".markdown",
    ".py",
    ".rst",
    ".text",
    ".toml",
    ".tsv",
    ".txt",
    ".yaml",
    ".yml",
}
DOCUMENT_SUFFIXES = {
    ".docx",
    ".pdf",
}
SPREADSHEET_SUFFIXES = {
    ".xls",
    ".xlsx",
}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | DOCUMENT_SUFFIXES | SPREADSHEET_SUFFIXES
DEFAULT_IGNORE = [".git/**", "**/.git/**", "__pycache__/**", "**/__pycache__/**", ".DS_Store", "**/.DS_Store"]
COLLECTION_MARKERS = (".pska-source.json", "pska-source.json")


@dataclass(slots=True)
class FilesScanReport:
    root: str
    connector_state: ConnectorState
    scanned: int = 0
    ingested: int = 0
    new_files: int = 0
    changed_files: int = 0
    unchanged_files: int = 0
    moved_files: int = 0
    missing_files: int = 0
    skipped: list[dict[str, Any]] = field(default_factory=list)
    failed: list[dict[str, Any]] = field(default_factory=list)
    source_item_ids: list[str] = field(default_factory=list)
    changes: list[dict[str, Any]] = field(default_factory=list)
    processing_config: dict[str, Any] = field(default_factory=dict)


def scan_files(
    store: KnowledgeStore,
    *,
    root: Path,
    owner_user_id: str = "user_primary",
    tenant_id: str = DEFAULT_TENANT_ID,
    space_id: str = "private_primary",
    visibility: Visibility = Visibility.PRIVATE,
    visible_team_ids: list[str] | None = None,
    ignore: list[str] | None = None,
    max_bytes: int = DEFAULT_FILES_MAX_BYTES,
    spreadsheet_max_rows_per_sheet: int = DEFAULT_SPREADSHEET_MAX_ROWS_PER_SHEET,
    spreadsheet_max_columns: int = DEFAULT_SPREADSHEET_MAX_COLUMNS,
    embedding_provider=None,
    processing_config: dict[str, Any] | None = None,
) -> FilesScanReport:
    root = root.expanduser().resolve()
    if not root.exists() or not root.is_dir():
        raise ValueError(f"Files connector root must be an existing directory: {root}")

    effective_processing_config = resolve_processing_config(processing_config)
    state = _connector_state(store, root=root, owner_user_id=owner_user_id, tenant_id=tenant_id)
    if not state.enabled:
        return FilesScanReport(
            root=str(root),
            connector_state=state,
            skipped=[{"root": str(root), "reason": "connector_disabled"}],
            processing_config=effective_processing_config,
        )

    ingest = IngestService(store, embedding_provider=embedding_provider, processing_config=effective_processing_config)
    report = FilesScanReport(root=str(root), connector_state=state, processing_config=effective_processing_config)
    latest_cursor = state.scan_cursor
    patterns = [*(ignore or []), *DEFAULT_IGNORE]
    root_key = str(root)
    state_config = dict(state.config or {})
    manifests_by_root = dict(state_config.get("files_manifests_by_root") or {})
    missing_by_root = dict(state_config.get("files_missing_by_root") or {})
    previous_manifest = dict(manifests_by_root.get(root_key) or {})
    if not previous_manifest:
        legacy_manifest = dict(state_config.get("files_manifest") or {})
        if _manifest_belongs_to_root(legacy_manifest, root):
            previous_manifest = legacy_manifest
    next_manifest: dict[str, dict[str, Any]] = {}
    previous_by_hash = _manifest_by_hash(previous_manifest)
    existing_source_item_ids = {item.source_item_id for item in store.list_source_items(tenant_id=tenant_id)}
    collection_roots = _find_collection_roots(root, ignore=patterns)

    for collection_root in collection_roots:
        relative = collection_root.relative_to(root).as_posix()
        try:
            classification = _classify_collection(collection_root, root, previous_manifest, previous_by_hash)
            previous_entry = previous_manifest.get(relative)
            if classification["status"] == "unchanged" and not _manifest_source_item_exists(previous_entry, existing_source_item_ids):
                classification = {"status": "new", "reason": "manifest_source_item_missing"}
            if classification["status"] == "unchanged":
                report.unchanged_files += 1
                next_manifest[relative] = {
                    **dict(previous_entry or {}),
                    "path": str(collection_root),
                    "relative_path": relative,
                    "content_hash": classification["content_hash"],
                    "collection": True,
                }
                report.changes.append({"path": str(collection_root), "status": "unchanged", "collection": True})
                continue
            item, documents = _ingest_collection(
                store,
                ingest,
                collection_root=collection_root,
                root=root,
                owner_user_id=owner_user_id,
                tenant_id=tenant_id,
                space_id=space_id,
                visibility=visibility,
                visible_team_ids=visible_team_ids or [],
                content_hash=classification["content_hash"],
                files=classification["files"],
            )
            report.scanned += len(documents)
            report.ingested += 1
            report.source_item_ids.append(item.source_item_id)
            _record_classification(report, classification)
            next_manifest[relative] = {
                "path": str(collection_root),
                "relative_path": relative,
                "content_hash": classification["content_hash"],
                "source_item_id": item.source_item_id,
                "collection": True,
                "document_count": len(documents),
                "files": [document.metadata.get("relative_path") for document in documents],
            }
            report.changes.append(
                {
                    "path": str(collection_root),
                    "status": classification["status"],
                    "source_item_id": item.source_item_id,
                    "content_hash": classification["content_hash"],
                    "collection": True,
                    "document_count": len(documents),
                }
            )
        except Exception as exc:  # noqa: BLE001 - per-collection failures should not abort a scan.
            report.failed.append({"path": str(collection_root), "error": f"{type(exc).__name__}: {exc}"})

    for path in _iter_candidate_files(root, ignore=patterns):
        if _is_within_collection(path, collection_roots):
            continue
        report.scanned += 1
        relative = path.relative_to(root).as_posix()
        try:
            stat = path.stat()
            if stat.st_size > max_bytes:
                report.skipped.append({"path": str(path), "reason": "file_too_large", "size_bytes": stat.st_size})
                continue
            suffix = path.suffix.lower()
            if suffix not in SUPPORTED_SUFFIXES:
                report.skipped.append({"path": str(path), "reason": "unsupported_suffix"})
                continue
            raw = path.read_bytes()
            file_content_hash = "sha256:" + hashlib.sha256(raw).hexdigest()
            classification = _classify_file(relative, file_content_hash, previous_manifest, previous_by_hash)
            if classification["status"] == "unchanged" and not _manifest_source_item_exists(
                previous_manifest.get(relative),
                existing_source_item_ids,
            ):
                classification = {"status": "new", "reason": "manifest_source_item_missing"}
            if classification["status"] == "unchanged":
                report.unchanged_files += 1
                next_manifest[relative] = {
                    **dict(previous_manifest.get(relative) or {}),
                    "path": str(path),
                    "relative_path": relative,
                    "content_hash": file_content_hash,
                    "size_bytes": stat.st_size,
                    "mtime_ns": str(int(stat.st_mtime_ns)),
                }
                report.changes.append({"path": str(path), "status": "unchanged"})
                latest_cursor = _max_scan_cursor(latest_cursor, str(int(stat.st_mtime_ns)))
                continue
            extracted = _extract_text(
                path,
                raw,
                spreadsheet_max_rows_per_sheet=spreadsheet_max_rows_per_sheet,
                spreadsheet_max_columns=spreadsheet_max_columns,
            )
            if not extracted["ok"]:
                report.skipped.append({"path": str(path), "reason": extracted["reason"], "detail": extracted.get("detail")})
                continue
            text = str(extracted["text"])
            if not text.strip():
                report.skipped.append({"path": str(path), "reason": "empty_text"})
                continue
            item = _ingest_file(
                ingest,
                path=path,
                root=root,
                relative=relative,
                raw=raw,
                text=text,
                owner_user_id=owner_user_id,
                tenant_id=tenant_id,
                space_id=space_id,
                visibility=visibility,
                visible_team_ids=visible_team_ids or [],
                stat=stat,
                content_hash=file_content_hash,
                extraction=extracted,
            )
            report.ingested += 1
            report.source_item_ids.append(item.source_item_id)
            _record_classification(report, classification)
            next_manifest[relative] = _manifest_entry(
                path=path,
                relative=relative,
                content_hash=file_content_hash,
                stat=stat,
                source_item_id=item.source_item_id,
            )
            change = {
                "path": str(path),
                "status": classification["status"],
                "source_item_id": item.source_item_id,
                "content_hash": file_content_hash,
            }
            if classification.get("previous_path"):
                change["previous_path"] = classification["previous_path"]
            report.changes.append(change)
            latest_cursor = _max_scan_cursor(latest_cursor, str(int(stat.st_mtime_ns)))
        except Exception as exc:  # noqa: BLE001 - per-file failures should not abort a scan.
            report.failed.append({"path": str(path), "error": f"{type(exc).__name__}: {exc}"})

    missing = _missing_manifest(previous_manifest, next_manifest)
    report.missing_files = len(missing)
    report.changes.extend(missing)
    status = "failed" if report.failed and not report.ingested else "succeeded"
    state.scan_cursor = latest_cursor
    state.sync_status = status
    manifests_by_root[root_key] = next_manifest
    missing_by_root[root_key] = missing
    state.config = {
        **state_config,
        "files_manifests_by_root": manifests_by_root,
        "files_missing_by_root": missing_by_root,
        "files_manifest": next_manifest,
        "files_missing": missing,
    }
    if report.failed:
        state.last_error = report.failed[0]["error"]
    else:
        state.last_error = None
    report.connector_state = store.upsert_connector_state(state)
    return report


def _extract_text(
    path: Path,
    raw: bytes,
    *,
    spreadsheet_max_rows_per_sheet: int = DEFAULT_SPREADSHEET_MAX_ROWS_PER_SHEET,
    spreadsheet_max_columns: int = DEFAULT_SPREADSHEET_MAX_COLUMNS,
) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return {"ok": True, "text": raw.decode("utf-8", errors="replace"), "extractor": "utf8"}
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore[import-not-found]
        except Exception as exc:  # noqa: BLE001 - optional dependency.
            return {"ok": False, "reason": "missing_dependency", "detail": f"pypdf required for PDF extraction: {type(exc).__name__}"}
        try:
            reader = PdfReader(path)
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:  # noqa: BLE001 - per-file failure should not abort scan.
            return {"ok": False, "reason": "extract_failed", "detail": f"{type(exc).__name__}: {exc}"}
        return {"ok": True, "text": text, "extractor": "pypdf"}
    if suffix == ".docx":
        try:
            from docx import Document  # type: ignore[import-not-found]
        except Exception as exc:  # noqa: BLE001 - optional dependency.
            return {"ok": False, "reason": "missing_dependency", "detail": f"python-docx required for DOCX extraction: {type(exc).__name__}"}
        try:
            document = Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception as exc:  # noqa: BLE001
            return {"ok": False, "reason": "extract_failed", "detail": f"{type(exc).__name__}: {exc}"}
        return {"ok": True, "text": text, "extractor": "python-docx"}
    if suffix == ".xlsx":
        return _extract_xlsx_text(
            path,
            raw,
            spreadsheet_max_rows_per_sheet=spreadsheet_max_rows_per_sheet,
            spreadsheet_max_columns=spreadsheet_max_columns,
        )
    if suffix == ".xls":
        return _extract_xls_text(
            path,
            spreadsheet_max_rows_per_sheet=spreadsheet_max_rows_per_sheet,
            spreadsheet_max_columns=spreadsheet_max_columns,
        )
    return {"ok": False, "reason": "unsupported_suffix"}


def extract_text_from_bytes(
    filename: str,
    raw: bytes,
    *,
    spreadsheet_max_rows_per_sheet: int = DEFAULT_SPREADSHEET_MAX_ROWS_PER_SHEET,
    spreadsheet_max_columns: int = DEFAULT_SPREADSHEET_MAX_COLUMNS,
) -> dict[str, Any]:
    """Extract uploaded file text with the same rules as folder sources."""
    safe_name = Path(filename or "upload.txt").name or "upload.txt"
    path = Path(safe_name)
    if path.suffix.lower() in TEXT_SUFFIXES:
        return _extract_text(
            path,
            raw,
            spreadsheet_max_rows_per_sheet=spreadsheet_max_rows_per_sheet,
            spreadsheet_max_columns=spreadsheet_max_columns,
        )
    with tempfile.TemporaryDirectory(prefix="pska_upload_extract_") as tmpdir:
        temp_path = Path(tmpdir) / safe_name
        temp_path.write_bytes(raw)
        return _extract_text(
            temp_path,
            raw,
            spreadsheet_max_rows_per_sheet=spreadsheet_max_rows_per_sheet,
            spreadsheet_max_columns=spreadsheet_max_columns,
        )


def _extract_xlsx_text(
    path: Path,
    raw: bytes,
    *,
    spreadsheet_max_rows_per_sheet: int,
    spreadsheet_max_columns: int,
) -> dict[str, Any]:
    row_limit = max(1, int(spreadsheet_max_rows_per_sheet))
    column_limit = max(1, int(spreadsheet_max_columns))
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            shared_strings = _xlsx_shared_strings(archive)
            sheets = _xlsx_sheet_paths(archive)
            rendered: list[str] = [f"# Workbook: {path.name}"]
            sheet_metadata: list[dict[str, Any]] = []
            for sheet_name, sheet_path in sheets:
                rows = _xlsx_sheet_rows(archive, sheet_path, shared_strings)
                rows = _trim_empty_rows(rows)
                max_columns = max((len(row) for row in rows), default=0)
                truncated_rows = len(rows) > row_limit
                truncated_columns = max_columns > column_limit
                rows = [row[:column_limit] for row in rows[:row_limit]]
                rendered.append(f"\n## Sheet: {sheet_name}\n")
                rendered.append(_rows_to_markdown_table(rows) if rows else "_Empty sheet._")
                sheet_metadata.append(
                    {
                        "name": sheet_name,
                        "rows": len(rows),
                        "columns": min(max_columns, column_limit),
                        "truncated_rows": truncated_rows,
                        "truncated_columns": truncated_columns,
                    }
                )
    except zipfile.BadZipFile as exc:
        return {"ok": False, "reason": "extract_failed", "detail": f"invalid xlsx zip: {exc}"}
    except Exception as exc:  # noqa: BLE001 - per-file failure should not abort scan.
        return {"ok": False, "reason": "extract_failed", "detail": f"{type(exc).__name__}: {exc}"}
    text = "\n\n".join(rendered)
    return {
        "ok": True,
        "text": text,
        "extractor": "xlsx-zip-xml",
        "metadata": {
            "sheet_count": len(sheet_metadata),
            "sheets": sheet_metadata,
            "row_limit_per_sheet": row_limit,
            "column_limit": column_limit,
        },
    }


def _extract_xls_text(
    path: Path,
    *,
    spreadsheet_max_rows_per_sheet: int,
    spreadsheet_max_columns: int,
) -> dict[str, Any]:
    row_limit = max(1, int(spreadsheet_max_rows_per_sheet))
    column_limit = max(1, int(spreadsheet_max_columns))
    try:
        import xlrd  # type: ignore[import-not-found]
    except Exception as exc:  # noqa: BLE001 - optional dependency.
        return {"ok": False, "reason": "missing_dependency", "detail": f"xlrd required for XLS extraction: {type(exc).__name__}"}
    try:
        workbook = xlrd.open_workbook(str(path))
        rendered: list[str] = [f"# Workbook: {path.name}"]
        sheet_metadata: list[dict[str, Any]] = []
        for sheet in workbook.sheets():
            row_count = min(sheet.nrows, row_limit)
            col_count = min(sheet.ncols, column_limit)
            rows = [
                [_cell_to_text(sheet.cell_value(row_index, col_index)) for col_index in range(col_count)]
                for row_index in range(row_count)
            ]
            rows = _trim_empty_rows(rows)
            rendered.append(f"\n## Sheet: {sheet.name}\n")
            rendered.append(_rows_to_markdown_table(rows) if rows else "_Empty sheet._")
            sheet_metadata.append(
                {
                    "name": sheet.name,
                    "rows": len(rows),
                    "columns": col_count,
                    "truncated_rows": sheet.nrows > row_limit,
                    "truncated_columns": sheet.ncols > column_limit,
                }
            )
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "reason": "extract_failed", "detail": f"{type(exc).__name__}: {exc}"}
    return {
        "ok": True,
        "text": "\n\n".join(rendered),
        "extractor": "xlrd",
        "metadata": {
            "sheet_count": len(sheet_metadata),
            "sheets": sheet_metadata,
            "row_limit_per_sheet": row_limit,
            "column_limit": column_limit,
        },
    }


def _xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        xml = archive.read("xl/sharedStrings.xml")
    except KeyError:
        return []
    root = ElementTree.fromstring(xml)
    values: list[str] = []
    for item in root:
        parts = [
            text_node.text or ""
            for text_node in item.iter()
            if text_node.tag.endswith("}t") or text_node.tag == "t"
        ]
        values.append("".join(parts))
    return values


def _xlsx_sheet_paths(archive: zipfile.ZipFile) -> list[tuple[str, str]]:
    workbook = ElementTree.fromstring(archive.read("xl/workbook.xml"))
    relationships = ElementTree.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
    rel_paths: dict[str, str] = {}
    for rel in relationships:
        rel_id = str(rel.attrib.get("Id") or "")
        target = str(rel.attrib.get("Target") or "")
        if not rel_id or not target:
            continue
        rel_paths[rel_id] = _xlsx_target_to_path(target)
    sheets: list[tuple[str, str]] = []
    for sheet in workbook.iter():
        if not sheet.tag.endswith("}sheet") and sheet.tag != "sheet":
            continue
        name = str(sheet.attrib.get("name") or f"Sheet {len(sheets) + 1}")
        rel_id = str(sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id") or "")
        sheet_path = rel_paths.get(rel_id)
        if sheet_path:
            sheets.append((name, sheet_path))
    return sheets


def _xlsx_target_to_path(target: str) -> str:
    if target.startswith("/"):
        return target.lstrip("/")
    if target.startswith("xl/"):
        return target
    return f"xl/{target}"


def _xlsx_sheet_rows(archive: zipfile.ZipFile, sheet_path: str, shared_strings: list[str]) -> list[list[str]]:
    sheet = ElementTree.fromstring(archive.read(sheet_path))
    rows: list[list[str]] = []
    for row in sheet.iter():
        if not row.tag.endswith("}row") and row.tag != "row":
            continue
        values: dict[int, str] = {}
        next_column = 1
        for cell in row:
            if not cell.tag.endswith("}c") and cell.tag != "c":
                continue
            cell_ref = str(cell.attrib.get("r") or "")
            column_index = _column_index_from_cell_ref(cell_ref) or next_column
            values[column_index] = _xlsx_cell_text(cell, shared_strings)
            next_column = column_index + 1
        max_column = max(values, default=0)
        rows.append([values.get(index, "") for index in range(1, max_column + 1)])
    return rows


def _xlsx_cell_text(cell: ElementTree.Element, shared_strings: list[str]) -> str:
    cell_type = cell.attrib.get("t")
    if cell_type == "inlineStr":
        return "".join(
            text_node.text or ""
            for text_node in cell.iter()
            if text_node.tag.endswith("}t") or text_node.tag == "t"
        )
    value_node = next((node for node in cell if node.tag.endswith("}v") or node.tag == "v"), None)
    raw_value = value_node.text if value_node is not None else ""
    if cell_type == "s":
        try:
            return shared_strings[int(raw_value or "0")]
        except Exception:  # noqa: BLE001
            return raw_value or ""
    if cell_type == "b":
        return "TRUE" if raw_value == "1" else "FALSE"
    return raw_value or ""


def _column_index_from_cell_ref(cell_ref: str) -> int | None:
    letters = "".join(char for char in cell_ref if char.isalpha())
    if not letters:
        return None
    index = 0
    for char in letters.upper():
        index = index * 26 + (ord(char) - ord("A") + 1)
    return index


def _trim_empty_rows(rows: list[list[str]]) -> list[list[str]]:
    return [row for row in rows if any(str(cell).strip() for cell in row)]


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    column_count = max(len(row) for row in rows)
    padded = [row + [""] * (column_count - len(row)) for row in rows]
    header = padded[0]
    body = padded[1:]
    if not any(str(cell).strip() for cell in header):
        header = [f"Column {index}" for index in range(1, column_count + 1)]
    lines = [
        "| " + " | ".join(_escape_table_cell(cell) for cell in header) + " |",
        "| " + " | ".join("---" for _ in range(column_count)) + " |",
    ]
    lines.extend("| " + " | ".join(_escape_table_cell(cell) for cell in row) + " |" for row in body)
    return "\n".join(lines)


def _escape_table_cell(value: Any) -> str:
    return str(value).replace("\n", " ").replace("\r", " ").replace("|", "\\|").strip()


def _cell_to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _connector_state(store: KnowledgeStore, *, root: Path, owner_user_id: str, tenant_id: str = DEFAULT_TENANT_ID) -> ConnectorState:
    state_id = connector_state_id("files", owner_user_id, tenant_id)
    try:
        state = store.get_connector_state(state_id)
        roots = state.permission_scope.setdefault("roots", [])
        if isinstance(roots, list) and str(root) not in roots:
            roots.append(str(root))
        return state
    except KeyError:
        state = connector_state_from_mapping(
            {
                "connector_id": "files",
                "owner_user_id": owner_user_id,
                "tenant_id": tenant_id,
                "enabled": True,
                "permission_scope": {"roots": [str(root)]},
                "config": {"default_ignore": DEFAULT_IGNORE},
            }
        )
        return store.upsert_connector_state(state)


def _iter_candidate_files(root: Path, *, ignore: list[str]) -> list[Path]:
    files = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        relative = path.relative_to(root).as_posix()
        if any(fnmatch(relative, pattern) or fnmatch(path.name, pattern) for pattern in ignore):
            continue
        files.append(path)
    return files


def _find_collection_roots(root: Path, *, ignore: list[str]) -> list[Path]:
    roots: list[Path] = []
    for marker_name in COLLECTION_MARKERS:
        for marker in sorted(root.rglob(marker_name)):
            if not marker.is_file():
                continue
            relative = marker.relative_to(root).as_posix()
            if any(fnmatch(relative, pattern) or fnmatch(marker.name, pattern) for pattern in ignore):
                continue
            roots.append(marker.parent)
    deduped = sorted(set(roots), key=lambda path: len(path.relative_to(root).parts))
    selected: list[Path] = []
    for candidate in deduped:
        if any(_path_is_relative_to(candidate, existing) for existing in selected):
            continue
        selected.append(candidate)
    return selected


def _is_within_collection(path: Path, collection_roots: list[Path]) -> bool:
    return any(_path_is_relative_to(path, root) for root in collection_roots)


def _path_is_relative_to(path: Path, root: Path) -> bool:
    try:
        path.relative_to(root)
    except ValueError:
        return False
    return True


def _collection_marker(collection_root: Path) -> tuple[Path, dict[str, Any]]:
    for marker_name in COLLECTION_MARKERS:
        marker = collection_root / marker_name
        if marker.exists():
            return marker, json.loads(marker.read_text(encoding="utf-8"))
    return collection_root / COLLECTION_MARKERS[0], {}


def _collection_files(collection_root: Path, marker: dict[str, Any]) -> list[Path]:
    include = [str(pattern) for pattern in marker.get("documents") or marker.get("include") or ["*.md"]]
    exclude = [str(pattern) for pattern in marker.get("exclude") or []]
    files: list[Path] = []
    for path in sorted(collection_root.rglob("*")):
        if not path.is_file() or path.name in COLLECTION_MARKERS:
            continue
        relative = path.relative_to(collection_root).as_posix()
        if any(fnmatch(relative, pattern) or fnmatch(path.name, pattern) for pattern in exclude):
            continue
        if not any(fnmatch(relative, pattern) or fnmatch(path.name, pattern) for pattern in include):
            continue
        files.append(path)
    return files


def _classify_collection(
    collection_root: Path,
    scan_root: Path,
    previous_manifest: dict[str, Any],
    previous_by_hash: dict[str, list[tuple[str, dict[str, Any]]]],
) -> dict[str, Any]:
    marker_path, marker = _collection_marker(collection_root)
    files = _collection_files(collection_root, marker)
    basis = [marker_path.read_text(encoding="utf-8") if marker_path.exists() else "{}"]
    for path in files:
        raw = path.read_bytes()
        basis.append(path.relative_to(collection_root).as_posix())
        basis.append(hashlib.sha256(raw).hexdigest())
    content_hash = "sha256:" + hashlib.sha256("\n".join(basis).encode("utf-8")).hexdigest()
    relative = collection_root.relative_to(scan_root).as_posix()
    classification = _classify_file(relative, content_hash, previous_manifest, previous_by_hash)
    return {**classification, "content_hash": content_hash, "files": files, "marker": marker}


def _ingest_collection(
    store: KnowledgeStore,
    ingest: IngestService,
    *,
    collection_root: Path,
    root: Path,
    owner_user_id: str,
    space_id: str,
    visibility: Visibility,
    visible_team_ids: list[str],
    content_hash: str,
    files: list[Path],
    tenant_id: str = DEFAULT_TENANT_ID,
) -> tuple[SourceItem, list[Document]]:
    marker_path, marker = _collection_marker(collection_root)
    title = str(marker.get("title") or collection_root.name)
    source_id = str(marker.get("source_id") or collection_root.relative_to(root).as_posix())
    source_item_id = f"src_{uuid5(NAMESPACE_URL, f'{tenant_id}:files-collection:{content_hash}:{source_id}').hex}"
    document_texts: list[tuple[Path, str]] = []
    for path in files:
        raw = path.read_bytes()
        extracted = _extract_text(path, raw)
        if extracted["ok"] and str(extracted["text"]).strip():
            document_texts.append((path, str(extracted["text"])))
    combined_text = "\n\n".join(
        f"# {path.relative_to(collection_root).as_posix()}\n\n{text}"
        for path, text in document_texts
    )
    item = SourceItem(
        source_item_id=source_item_id,
        source_channel="files",
        record_type="file_collection",
        source_id=source_id,
        owner_user_id=owner_user_id,
        space_id=space_id,
        visibility=visibility,
        visible_team_ids=visible_team_ids,
        title=title,
        url=collection_root.as_uri(),
        content_text=combined_text,
        content_hash=content_hash,
        tenant_id=tenant_id,
        metadata={
            "schema_version": "pska.files.collection.v1",
            "collection": True,
            "marker_path": str(marker_path),
            "source_root": str(root),
            "relative_path": collection_root.relative_to(root).as_posix(),
            "document_count": len(document_texts),
            "extra": dict(marker.get("extra") or {}),
        },
    )
    stored = store.upsert_source_item(item)
    documents: list[Document] = []
    chunks: list[Chunk] = []
    for path, text in document_texts:
        relative = path.relative_to(collection_root).as_posix()
        document_id = f"doc_{source_item_id[4:]}_{uuid5(NAMESPACE_URL, source_id + ':' + relative).hex[:12]}"
        document = Document(
            document_id=document_id,
            source_item_id=stored.source_item_id,
            owner_user_id=stored.owner_user_id,
            space_id=stored.space_id,
            visibility=stored.visibility,
            visible_team_ids=stored.visible_team_ids,
            title=relative,
            body=text,
            metadata={
                "url": path.as_uri(),
                "path": str(path),
                "relative_path": relative,
                "collection_root": str(collection_root),
                "collection_source_id": source_id,
            },
            tenant_id=stored.tenant_id,
        )
        documents.append(document)
        chunk_spans = ingest._chunk_spans(text)  # noqa: SLF001 - collection ingest shares the existing chunking policy.
        embedding_texts = [span.embedding_text() for span in chunk_spans]
        chunk_embeddings = ingest.embedding_provider.embed_texts(embedding_texts) if ingest.embedding_provider else [None] * len(chunk_spans)
        chunk_prefix = f"chk_{source_item_id[4:]}_{uuid5(NAMESPACE_URL, source_id + ':' + relative).hex[:12]}"
        for ordinal, (span, embedding) in enumerate(zip(chunk_spans, chunk_embeddings)):
            chunks.append(
                Chunk(
                    chunk_id=f"{chunk_prefix}_{ordinal}",
                    document_id=document.document_id,
                    source_item_id=stored.source_item_id,
                    owner_user_id=stored.owner_user_id,
                    space_id=stored.space_id,
                    visibility=stored.visibility,
                    visible_team_ids=stored.visible_team_ids,
                    text=span.text,
                    ordinal=ordinal,
                    embedding=embedding,
                    metadata={
                        "embedding_provider": ingest.embedding_provider.provider_name if ingest.embedding_provider else None,
                        "embedding_model": ingest.embedding_provider.model_name if ingest.embedding_provider else None,
                        "chunk_size": ingest.chunk_size,
                        "chunk_overlap": ingest.chunk_overlap,
                        "chunk_strategy": span.strategy,
                        "start": span.start,
                        "end": span.end,
                        "context_header": span.context_header,
                        "collection": True,
                        "relative_path": relative,
                    },
                    tenant_id=stored.tenant_id,
                )
            )
    store.replace_source_documents(stored.source_item_id, documents, chunks)
    OfflineIndexService(store, embedding_provider=ingest.embedding_provider).mark_source_dirty(
        stored,
        chunks,
        reason="source_collection_ingested",
    )
    return stored, documents


def _ingest_file(
    ingest: IngestService,
    *,
    path: Path,
    root: Path,
    relative: str,
    raw: bytes,
    text: str,
    owner_user_id: str,
    space_id: str,
    visibility: Visibility,
    visible_team_ids: list[str],
    tenant_id: str,
    stat,
    content_hash: str,
    extraction: dict[str, Any],
) -> SourceItem:
    mime_type = mimetypes.guess_type(path.name)[0] or "text/plain"
    extraction_metadata = {
        "extractor": extraction.get("extractor"),
        **dict(extraction.get("metadata") or {}),
    }
    payload = connector_record_to_payload(
        {
            "connector_id": "files",
            "external_id": str(path),
            "source_uri": path.as_uri(),
            "record_type": "file",
            "title": path.name,
            "body": text,
            "owner_user_id": owner_user_id,
            "tenant_id": tenant_id,
            "space_id": space_id,
            "visibility": visibility.value,
            "visible_team_ids": visible_team_ids,
            "updated_at": str(int(stat.st_mtime_ns)),
            "captured_at": None,
            "artifacts": {"path": str(path)},
            "permission_metadata": {"root": str(root), "relative_path": relative, "read_scope": "explicit_directory"},
            "scan_cursor": str(int(stat.st_mtime_ns)),
            "content_hash": content_hash,
            "metadata": {"mime_type": mime_type, "size_bytes": stat.st_size, "extraction": extraction_metadata},
        }
    )
    return ingest.ingest_channel_payload(payload)


def _manifest_by_hash(manifest: dict[str, Any]) -> dict[str, list[tuple[str, dict[str, Any]]]]:
    by_hash: dict[str, list[tuple[str, dict[str, Any]]]] = {}
    for relative, entry in manifest.items():
        if not isinstance(entry, dict):
            continue
        content_hash = entry.get("content_hash")
        if content_hash:
            by_hash.setdefault(str(content_hash), []).append((relative, entry))
    return by_hash


def _classify_file(
    relative: str,
    content_hash: str,
    previous_manifest: dict[str, Any],
    previous_by_hash: dict[str, list[tuple[str, dict[str, Any]]]],
) -> dict[str, Any]:
    previous = previous_manifest.get(relative)
    if isinstance(previous, dict):
        if previous.get("content_hash") == content_hash:
            return {"status": "unchanged"}
        return {"status": "changed", "previous_content_hash": previous.get("content_hash")}
    for previous_relative, entry in previous_by_hash.get(content_hash, []):
        if previous_relative != relative:
            return {"status": "moved", "previous_path": entry.get("path"), "previous_relative_path": previous_relative}
    return {"status": "new"}


def _record_classification(report: FilesScanReport, classification: dict[str, Any]) -> None:
    status = classification["status"]
    if status == "new":
        report.new_files += 1
    elif status == "changed":
        report.changed_files += 1
    elif status == "moved":
        report.moved_files += 1


def _manifest_source_item_exists(entry: Any, existing_source_item_ids: set[str]) -> bool:
    if not isinstance(entry, dict):
        return False
    source_item_id = entry.get("source_item_id")
    return bool(source_item_id and str(source_item_id) in existing_source_item_ids)


def _manifest_entry(*, path: Path, relative: str, content_hash: str, stat, source_item_id: str) -> dict[str, Any]:
    return {
        "path": str(path),
        "relative_path": relative,
        "content_hash": content_hash,
        "source_item_id": source_item_id,
        "size_bytes": stat.st_size,
        "mtime_ns": str(int(stat.st_mtime_ns)),
    }


def _missing_manifest(previous_manifest: dict[str, Any], next_manifest: dict[str, Any]) -> list[dict[str, Any]]:
    missing = []
    next_hashes = {entry.get("content_hash") for entry in next_manifest.values() if isinstance(entry, dict)}
    for relative, entry in sorted(previous_manifest.items()):
        if relative in next_manifest:
            continue
        if not isinstance(entry, dict):
            continue
        if entry.get("content_hash") in next_hashes:
            continue
        missing.append(
            {
                "status": "missing",
                "path": entry.get("path"),
                "relative_path": relative,
                "source_item_id": entry.get("source_item_id"),
                "content_hash": entry.get("content_hash"),
            }
        )
    return missing


def _manifest_belongs_to_root(manifest: dict[str, dict[str, Any]], root: Path) -> bool:
    if not manifest:
        return False
    resolved_root = root.expanduser().resolve()
    for entry in manifest.values():
        if not isinstance(entry, dict):
            return False
        path = entry.get("path")
        if not path:
            return False
        try:
            Path(str(path)).expanduser().resolve().relative_to(resolved_root)
        except (OSError, ValueError):
            return False
    return True


def _max_scan_cursor(current: str | None, candidate: str) -> str:
    if not current:
        return candidate
    try:
        return str(max(int(current), int(candidate)))
    except ValueError:
        return candidate
